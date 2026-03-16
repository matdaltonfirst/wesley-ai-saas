import os
import re
import json
import uuid
import time
import string
import secrets
import logging
import threading
from collections import defaultdict
from datetime import datetime, timedelta
from functools import wraps
from pathlib import Path

import click
import resend
import stripe
from flask import Flask, request, jsonify, render_template, redirect, url_for, make_response, send_from_directory, session, abort
from werkzeug.middleware.proxy_fix import ProxyFix
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from dotenv import load_dotenv
from google import genai
from google.genai import types
from sqlalchemy import text, inspect as sa_inspect
from sqlalchemy.orm import joinedload
import pdfplumber
from docx import Document as DocxDocument
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger

from models import db, User, Church, Document, SystemPrompt, CrawledPage, Conversation, Message, WidgetConversation, WidgetMessage, Invite

load_dotenv()

# ── Logging ────────────────────────────────────────────────────────────────────

log = logging.getLogger("wesley")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)


# ── In-memory rate limiter (per-IP, no extra dependency) ─────────────────────

class _RateLimiter:
    """Simple sliding-window rate limiter keyed by IP address."""

    def __init__(self, max_requests: int = 30, window_seconds: int = 60):
        self.max_requests = max_requests
        self.window = window_seconds
        self._hits: dict[str, list[float]] = defaultdict(list)
        self._lock = threading.Lock()

    def is_limited(self, key: str) -> bool:
        now = time.monotonic()
        with self._lock:
            timestamps = self._hits[key]
            # Prune old entries
            self._hits[key] = [t for t in timestamps if now - t < self.window]
            if len(self._hits[key]) >= self.max_requests:
                return True
            self._hits[key].append(now)
            return False

# Widget chat: 30 requests/minute per IP
_widget_chat_limiter = _RateLimiter(max_requests=30, window_seconds=60)
# Widget branding: 60 requests/minute per IP (lightweight)
_widget_branding_limiter = _RateLimiter(max_requests=60, window_seconds=60)


# ── In-memory document chunk cache ───────────────────────────────────────────
# Keyed by (doc_id, uploaded_at_iso). Since uploaded_at never changes for a
# given document row, entries are permanently valid until the document is
# deleted (at which point we proactively evict). Capped at 200 entries to
# bound memory usage (~50–100 KB per typical document = ~10–20 MB max).

_doc_cache: dict[tuple, list[dict]] = {}
_doc_cache_lock = threading.Lock()
_DOC_CACHE_MAX = 200


def _parse_doc_chunks(doc, filepath: "Path") -> list[dict]:
    """Return parsed chunks for a document, reading from cache when possible."""
    cache_key = (doc.id, doc.uploaded_at.isoformat())
    with _doc_cache_lock:
        if cache_key in _doc_cache:
            return _doc_cache[cache_key]

    suffix = filepath.suffix.lower()
    if suffix == ".pdf":
        chunks = read_pdf(filepath, doc.original_name)
    elif suffix == ".docx":
        chunks = read_docx(filepath, doc.original_name)
    else:
        chunks = []

    with _doc_cache_lock:
        if len(_doc_cache) >= _DOC_CACHE_MAX:
            # Evict oldest half when cap is hit
            evict_keys = list(_doc_cache.keys())[: _DOC_CACHE_MAX // 2]
            for k in evict_keys:
                del _doc_cache[k]
        _doc_cache[cache_key] = chunks

    return chunks


def _evict_doc_cache(doc_id: int, uploaded_at) -> None:
    """Remove a document's parsed chunks from the cache after deletion."""
    cache_key = (doc_id, uploaded_at.isoformat())
    with _doc_cache_lock:
        _doc_cache.pop(cache_key, None)


# ── CSRF protection (for HTML form POSTs) ─────────────────────────────────────

def _csrf_token() -> str:
    """Return (and lazily create) a per-session CSRF token."""
    if "csrf_token" not in session:
        session["csrf_token"] = secrets.token_hex(32)
    return session["csrf_token"]


def _validate_csrf() -> None:
    """Abort 403 if the submitted CSRF token doesn't match the session token."""
    token = request.form.get("csrf_token") or request.headers.get("X-CSRFToken", "")
    if not token or not secrets.compare_digest(token, session.get("csrf_token", "")):
        abort(403)


stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "")
resend.api_key = os.getenv("RESEND_API_KEY", "")

# ── Paths ─────────────────────────────────────────────────────────────────────

DATA_DIR = Path(os.getenv("DATA_DIR", "data")).resolve()
DATA_DIR.mkdir(parents=True, exist_ok=True)

UPLOADS_DIR = DATA_DIR / "uploads"
UPLOADS_DIR.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_UPLOAD_MB = 32

# ── App setup ─────────────────────────────────────────────────────────────────

app = Flask(__name__, static_folder="static", template_folder="templates")
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

_secret = os.getenv("SECRET_KEY", "")
if not _secret:
    _secret = secrets.token_hex(32)
    print("WARNING: SECRET_KEY is not set. Generated a random key — sessions will not persist across restarts.")
app.config["SECRET_KEY"] = _secret
app.config["SQLALCHEMY_DATABASE_URI"] = f"sqlite:///{DATA_DIR / 'wesley.db'}"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_MB * 1024 * 1024

db.init_app(app)

# Make csrf_token() available in all Jinja2 templates
app.jinja_env.globals["csrf_token"] = _csrf_token

login_manager = LoginManager(app)
login_manager.login_view = "login_page"


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


@login_manager.unauthorized_handler
def unauthorized():
    if request.path.startswith("/api/"):
        return jsonify({"error": "Authentication required."}), 401
    return redirect(url_for("login_page"))


# ── Platform constants (override via environment variables) ───────────────────

APP_URL       = os.getenv("APP_URL",       "https://app.wesleyai.co")
FROM_EMAIL    = os.getenv("FROM_EMAIL",    "Wesley AI <noreply@wesleyai.co>")
SUPPORT_EMAIL = os.getenv("SUPPORT_EMAIL", "info@wesleyai.co")
GEMINI_MODEL  = os.getenv("GEMINI_MODEL",  "gemini-2.5-flash")

# ── Branding defaults (single source of truth shared with the JS via API) ─────

DEFAULT_BOT_NAME = "Wesley"
DEFAULT_WELCOME  = "How can I help you today?"
DEFAULT_COLOR    = "#0a3d3d"
DEFAULT_SUBTITLE = "Ask me anything about our church"
DEFAULT_STARTERS = [
    "What is our volunteer policy?",
    "Help me draft a Sunday bulletin",
    "What events are coming up?",
    "Write a prayer for our newsletter",
]

# ── Super admin + default prompt ──────────────────────────────────────────────

SUPER_ADMIN_EMAIL = os.getenv("SUPER_ADMIN_EMAIL", "info@wesleyai.co")

DEFAULT_SYSTEM_PROMPT = (
    "You are Wesley, a helpful AI assistant for United Methodist churches. "
    "You are grounded in Wesleyan theology and United Methodist doctrine. "
    "You speak with warmth, grace, and pastoral care. "
    "You never contradict UMC doctrine. "
    "For deep theological or personal questions you always encourage the user "
    "to speak with their pastor."
)


def is_super_admin() -> bool:
    return current_user.is_authenticated and current_user.email == SUPER_ADMIN_EMAIL


def _build_branding_dict(church) -> dict:
    """Return the standard branding JSON dict for a Church record.

    Single source of truth for the branding payload — used by both the
    authenticated ``/api/church/branding`` endpoint and the public
    ``/api/widget/branding`` endpoint so the two can never diverge.
    """
    try:
        sugs = json.loads(church.starter_questions) if church.starter_questions else []
    except (ValueError, TypeError):
        sugs = []
    return {
        "bot_name":          church.bot_name       or DEFAULT_BOT_NAME,
        "bot_subtitle":      church.bot_subtitle    or DEFAULT_SUBTITLE,
        "welcome_message":   church.welcome_message or DEFAULT_WELCOME,
        "primary_color":     church.primary_color   or DEFAULT_COLOR,
        "church_city":       church.church_city     or "",
        "starter_questions": sugs,
    }


def _build_system_prompt(church, widget: bool = False) -> str:
    """Assemble the full Gemini system instruction for a given church.

    Both ``chat()`` (staff) and ``widget_chat()`` (public) call this helper
    so the church-identity block can never drift between the two code paths.
    Pass ``widget=True`` to add date-awareness, source-confidentiality, and
    a plain-text-only instruction (visitor-facing responses must not contain
    markdown symbols that the widget cannot render).
    """
    today_str = datetime.utcnow().strftime("%A, %B %-d, %Y")
    prompt_row = SystemPrompt.query.get(1)
    base = f"Today's date is {today_str}.\n\n" + (prompt_row.content if prompt_row else DEFAULT_SYSTEM_PROMPT)

    ctx = f"\n\nYou are installed at {church.name}"
    if church.church_city:
        ctx += f", located in {church.church_city}"
    ctx += f". Your name is {church.bot_name or DEFAULT_BOT_NAME}."

    if not widget:
        return base + ctx

    addendum = (
        "\n\nWhen answering questions about schedules, events, menus, or anything "
        "time-sensitive, use today's date to give a specific, direct answer — "
        "do not list every option when only today's is relevant."
        "\n\nIMPORTANT: Never mention that you are referencing a document, file, "
        "or uploaded file of any kind. Never reveal or repeat file names (including "
        ".pdf and .docx filenames). Answer naturally and directly, as if you simply "
        "know the information."
        "\n\nRespond in plain text only. Do not use markdown formatting such as "
        "headings (##), bullet points (-), bold (**text**), italic (*text*), "
        "or any other markdown syntax. Write in natural, conversational sentences."
    )
    return base + ctx + addendum


with app.app_context():
    db.create_all()
    log.info("db.create_all() completed — all tables present.")

    # Inline migration: add Phase 2 columns to existing churches table if absent
    insp = sa_inspect(db.engine)
    existing_cols = {c["name"] for c in insp.get_columns("churches")}
    with db.engine.connect() as conn:
        if "website_url" not in existing_cols:
            conn.execute(text("ALTER TABLE churches ADD COLUMN website_url VARCHAR(500)"))
            conn.commit()
            log.info("Migration: added churches.website_url")
        if "last_crawled_at" not in existing_cols:
            conn.execute(text("ALTER TABLE churches ADD COLUMN last_crawled_at DATETIME"))
            conn.commit()
            log.info("Migration: added churches.last_crawled_at")

    # Inline migration: branding columns added to churches table
    insp2 = sa_inspect(db.engine)
    existing_cols2 = {c["name"] for c in insp2.get_columns("churches")}
    with db.engine.connect() as conn2:
        if "bot_name" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN bot_name VARCHAR(100) NOT NULL DEFAULT 'Wesley'"))
            conn2.commit()
            log.info("Migration: added churches.bot_name")
        if "welcome_message" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN welcome_message VARCHAR(500) NOT NULL DEFAULT 'How can I help you today?'"))
            conn2.commit()
            log.info("Migration: added churches.welcome_message")
        if "primary_color" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN primary_color VARCHAR(7) NOT NULL DEFAULT '#0a3d3d'"))
            conn2.commit()
            log.info("Migration: added churches.primary_color")
        if "church_city" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN church_city VARCHAR(200)"))
            conn2.commit()
            log.info("Migration: added churches.church_city")
        if "onboarding_complete" not in existing_cols2:
            # DEFAULT 1 so all *existing* churches are treated as already onboarded;
            # new churches created via SQLAlchemy use the model default (False).
            conn2.execute(text("ALTER TABLE churches ADD COLUMN onboarding_complete BOOLEAN NOT NULL DEFAULT 1"))
            conn2.commit()
            log.info("Migration: added churches.onboarding_complete")
        if "trial_ends_at" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN trial_ends_at DATETIME"))
            conn2.commit()
            log.info("Migration: added churches.trial_ends_at")
        if "stripe_subscription_id" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN stripe_subscription_id VARCHAR(200)"))
            conn2.commit()
            log.info("Migration: added churches.stripe_subscription_id")
        if "billing_exempt" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN billing_exempt BOOLEAN NOT NULL DEFAULT 0"))
            conn2.commit()
            log.info("Migration: added churches.billing_exempt")
        if "plan" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN plan VARCHAR(20) NOT NULL DEFAULT 'founders'"))
            conn2.commit()
            log.info("Migration: added churches.plan")
        if "stripe_customer_id" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN stripe_customer_id VARCHAR(200)"))
            conn2.commit()
            log.info("Migration: added churches.stripe_customer_id")
        if "trial_reminder_sent" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN trial_reminder_sent BOOLEAN NOT NULL DEFAULT 0"))
            conn2.commit()
            log.info("Migration: added churches.trial_reminder_sent")
        if "starter_questions" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN starter_questions TEXT"))
            conn2.commit()
            log.info("Migration: added churches.starter_questions")
        if "bot_subtitle" not in existing_cols2:
            conn2.execute(text("ALTER TABLE churches ADD COLUMN bot_subtitle VARCHAR(200)"))
            conn2.commit()
            log.info("Migration: added churches.bot_subtitle")

    # Backfill trial_ends_at for any existing churches that don't have one yet.
    # Gives them a 14-day grace window from the date this migration runs.
    with db.engine.connect() as conn3:
        trial_cutoff = datetime.utcnow() + timedelta(days=14)
        result = conn3.execute(
            text("UPDATE churches SET trial_ends_at = :ts WHERE trial_ends_at IS NULL"),
            {"ts": trial_cutoff},
        )
        conn3.commit()
        if result.rowcount:
            log.info("Migration: set trial_ends_at for %d existing church(es)", result.rowcount)

    # Inline migration: add visibility column to documents table
    insp_docs = sa_inspect(db.engine)
    existing_doc_cols = {c["name"] for c in insp_docs.get_columns("documents")}
    with db.engine.connect() as conn_d:
        if "visibility" not in existing_doc_cols:
            conn_d.execute(text(
                "ALTER TABLE documents ADD COLUMN visibility VARCHAR(20) NOT NULL DEFAULT 'staff_only'"
            ))
            conn_d.commit()
            log.info("Migration: added documents.visibility (default 'staff_only')")

    # Inline migration: add password-reset + role columns to users table
    insp_users = sa_inspect(db.engine)
    existing_user_cols = {c["name"] for c in insp_users.get_columns("users")}
    with db.engine.connect() as conn_u:
        if "reset_token" not in existing_user_cols:
            conn_u.execute(text("ALTER TABLE users ADD COLUMN reset_token VARCHAR(100)"))
            conn_u.commit()
            log.info("Migration: added users.reset_token")
        if "reset_token_expires" not in existing_user_cols:
            conn_u.execute(text("ALTER TABLE users ADD COLUMN reset_token_expires DATETIME"))
            conn_u.commit()
            log.info("Migration: added users.reset_token_expires")
        if "role" not in existing_user_cols:
            # DEFAULT 'admin' so all existing users become admins
            conn_u.execute(text("ALTER TABLE users ADD COLUMN role VARCHAR(20) NOT NULL DEFAULT 'admin'"))
            conn_u.commit()
            log.info("Migration: added users.role")

    # Seed the master system prompt on first run (id=1 is the single canonical row)
    if not SystemPrompt.query.get(1):
        db.session.add(SystemPrompt(id=1, content=DEFAULT_SYSTEM_PROMPT))
        db.session.commit()
        log.info("System prompt seeded with default.")

# ── Flask CLI commands ────────────────────────────────────────────────────────


@app.cli.command("init-db")
def init_db_command():
    """Explicitly create all database tables. Safe to run on an existing DB."""
    db.create_all()
    click.echo("init-db: all tables created (or already exist).")
    # Report which tables are present
    from sqlalchemy import inspect as sa_inspect2
    tables = sa_inspect2(db.engine).get_table_names()
    click.echo(f"init-db: tables in DB → {', '.join(sorted(tables))}")


# ── API key validation ────────────────────────────────────────────────────────

_api_key = os.getenv("GEMINI_API_KEY")
if not _api_key:
    log.warning("GEMINI_API_KEY is not set. Copy .env.example to .env and add your key.")
else:
    log.info("Gemini API key loaded (%s…)", _api_key[:8])

if not os.getenv("STRIPE_ANNUAL_PRICE_ID"):
    log.warning("STRIPE_ANNUAL_PRICE_ID is not set. Annual billing will not work.")

if not os.getenv("RESEND_API_KEY"):
    log.warning("RESEND_API_KEY is not set. Password reset emails will not be sent.")

# ── Constants ─────────────────────────────────────────────────────────────────

STOP_WORDS = {
    "a", "an", "the", "is", "are", "was", "were", "be", "been", "being",
    "have", "has", "had", "do", "does", "did", "will", "would", "could",
    "should", "may", "might", "shall", "can", "need", "dare", "ought",
    "used", "to", "of", "in", "on", "at", "by", "for", "with", "about",
    "against", "between", "into", "through", "during", "before", "after",
    "above", "below", "from", "up", "down", "out", "off", "over", "under",
    "again", "further", "then", "once", "and", "but", "or", "nor", "so",
    "yet", "both", "either", "neither", "not", "only", "own", "same",
    "than", "too", "very", "just", "this", "that", "these", "those",
    "i", "me", "my", "we", "our", "you", "your", "he", "she", "it",
    "his", "her", "its", "they", "them", "their", "what", "which", "who",
    "whom", "how", "when", "where", "why", "all", "each", "every", "any",
    "no", "more", "most", "other", "some", "such", "if", "as",
}

# ── Document processing ───────────────────────────────────────────────────────


def get_church_dir(church_id: int) -> Path:
    d = UPLOADS_DIR / str(church_id)
    d.mkdir(exist_ok=True)
    return d


def read_pdf(filepath: Path, display_name: str) -> list[dict]:
    chunks = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text_content = page.extract_text()
                if text_content and text_content.strip():
                    chunks.append({
                        "content": text_content.strip(),
                        "source": display_name,
                        "location": f"Page {page_num}",
                    })
    except Exception as e:
        log.error("Error reading PDF %s: %s", filepath.name, e)
    return chunks


def read_docx(filepath: Path, display_name: str) -> list[dict]:
    chunks = []
    try:
        doc = DocxDocument(filepath)
        current_heading = "Document"
        current_text: list[str] = []
        current_length = 0

        def flush_chunk():
            nonlocal current_text, current_length
            text_content = " ".join(current_text).strip()
            if text_content:
                chunks.append({
                    "content": text_content,
                    "source": display_name,
                    "location": current_heading,
                })
            current_text = []
            current_length = 0

        for para in doc.paragraphs:
            style = para.style.name.lower()
            para_text = para.text.strip()
            if not para_text:
                continue
            if "heading" in style:
                flush_chunk()
                current_heading = para_text
            else:
                if current_length + len(para_text) > 500 and current_text:
                    flush_chunk()
                current_text.append(para_text)
                current_length += len(para_text)

        flush_chunk()
    except Exception as e:
        log.error("Error reading DOCX %s: %s", filepath.name, e)
    return chunks


def load_church_documents(church_id: int) -> list[dict]:
    """Load and parse all documents for a church (staff chat — no visibility filter)."""
    docs = Document.query.filter_by(church_id=church_id).all()
    church_dir = get_church_dir(church_id)
    all_chunks = []
    for doc in docs:
        filepath = church_dir / doc.filename
        if not filepath.exists():
            continue
        all_chunks.extend(_parse_doc_chunks(doc, filepath))
    return all_chunks


def load_chatbot_documents(church_id: int) -> list[dict]:
    """Load and parse only documents marked staff_and_chatbot (widget chat)."""
    docs = Document.query.filter_by(church_id=church_id, visibility="staff_and_chatbot").all()
    church_dir = get_church_dir(church_id)
    all_chunks = []
    for doc in docs:
        filepath = church_dir / doc.filename
        if not filepath.exists():
            continue
        all_chunks.extend(_parse_doc_chunks(doc, filepath))
    return all_chunks


def load_church_web_content(church_id: int) -> list[dict]:
    """Return keyword-scoreable chunks from a church's crawled web pages."""
    pages = CrawledPage.query.filter_by(church_id=church_id).all()
    chunks = []
    for page in pages:
        if page.content and page.content.strip():
            chunks.append({
                "content": page.content,
                "source": page.title or page.url,
                "location": page.url,
            })
    return chunks


# ── Relevance scoring ─────────────────────────────────────────────────────────


def extract_keywords(query: str) -> list[str]:
    words = re.findall(r"\b[a-zA-Z]{3,}\b", query.lower())
    return [w for w in words if w not in STOP_WORDS]


def score_chunk(chunk: dict, keywords: list[str]) -> int:
    content_lower = chunk["content"].lower()
    return sum(content_lower.count(kw) for kw in keywords)


def find_relevant_chunks(query: str, chunks: list[dict], top_n: int = 8) -> list[tuple[int, dict]]:
    keywords = extract_keywords(query)
    if not keywords:
        return []
    scored = [(score_chunk(c, keywords), c) for c in chunks]
    scored.sort(key=lambda x: x[0], reverse=True)
    return [(score, c) for score, c in scored[:top_n] if score > 0]


def build_context_block(scored_chunks: list[tuple[int, dict]]) -> str:
    lines = []
    for _, chunk in scored_chunks:
        lines.append(f"[From: {chunk['source']}, {chunk['location']}]")
        lines.append(chunk["content"])
        lines.append("")
    return "\n".join(lines)


# ── Gemini call ───────────────────────────────────────────────────────────────


def _friendly_gemini_error(exc: Exception) -> tuple[str, int]:
    msg = str(exc).lower()
    if "429" in msg or "quota" in msg or "rate" in msg or "exhausted" in msg:
        return ("The AI service is temporarily over its request limit. Please wait and try again.", 429)
    if "401" in msg or "403" in msg or "api_key" in msg:
        return ("API key error — please check that GEMINI_API_KEY is configured correctly.", 401)
    if "404" in msg or "not found" in msg:
        return ("The AI model could not be found. Please check the model name.", 404)
    if "503" in msg or "unavailable" in msg:
        return ("The AI service is temporarily unavailable. Please try again.", 503)
    return (f"AI error: {exc}", 502)


def call_gemini(question: str, context: str, history: list[dict], system_instruction: str) -> str:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY is not set. Add it to your .env file.")

    client = genai.Client(api_key=api_key)

    contents: list[types.Content] = []
    for msg in history:
        role = "user" if msg["role"] == "user" else "model"
        contents.append(types.Content(role=role, parts=[types.Part(text=msg["content"])]))

    current_text = (
        f"[Relevant church information:]\n{context}\n---\n{question}"
        if context.strip()
        else question
    )
    contents.append(types.Content(role="user", parts=[types.Part(text=current_text)]))

    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=contents,
        config=types.GenerateContentConfig(system_instruction=system_instruction),
    )
    return response.text


# ── Nightly crawl scheduler ───────────────────────────────────────────────────


def nightly_crawl_job():
    """Re-crawl all churches that have a website URL configured. Runs at 2am daily."""
    with app.app_context():
        from crawler import crawl_church_website
        churches = Church.query.filter(Church.website_url.isnot(None)).all()
        log.info("Nightly crawl: found %d church(es) to crawl.", len(churches))
        for church in churches:
            if not church.website_url:
                continue
            try:
                result = crawl_church_website(church.id, church.website_url)
                log.info("Nightly crawl church_id=%d (%s): %s", church.id, church.name, result)
            except Exception as exc:
                log.error("Nightly crawl error church_id=%d: %s", church.id, exc)


def nightly_cleanup_job():
    """Delete conversations (and their messages) last updated more than 14 days ago."""
    with app.app_context():
        cutoff = datetime.utcnow() - timedelta(days=14)
        old_convs = Conversation.query.filter(Conversation.updated_at < cutoff).all()
        count = len(old_convs)
        for conv in old_convs:
            db.session.delete(conv)
        db.session.commit()
        log.info("Nightly cleanup: deleted %d staff conversation(s) older than 14 days.", count)


def nightly_widget_cleanup_job():
    """Delete widget conversations (and their messages) older than 30 days."""
    with app.app_context():
        cutoff = datetime.utcnow() - timedelta(days=30)
        old = WidgetConversation.query.filter(WidgetConversation.updated_at < cutoff).all()
        count = len(old)
        for wconv in old:
            db.session.delete(wconv)
        db.session.commit()
        log.info("Nightly widget cleanup: deleted %d widget conversation(s) older than 30 days.", count)


def trial_reminder_job():
    """Daily 9 AM job: email churches whose trial ends in 3–5 days (once only)."""
    with app.app_context():
        now  = datetime.utcnow()
        low  = now + timedelta(days=3)
        high = now + timedelta(days=5)
        churches = Church.query.filter(
            Church.trial_ends_at >= low,
            Church.trial_ends_at <= high,
            Church.trial_reminder_sent == False,  # noqa: E712
            Church.stripe_subscription_id == None,  # noqa: E711
            Church.billing_exempt == False,  # noqa: E712
        ).all()
        sent = 0
        for church in churches:
            first_user = User.query.filter_by(church_id=church.id).order_by(User.id).first()
            if first_user:
                _send_trial_expiring_email(first_user.email, church.name, church.trial_ends_at)
            church.trial_reminder_sent = True
            sent += 1
        if churches:
            db.session.commit()
        log.info("Trial reminder job: sent %d reminder(s).", sent)


def invite_cleanup_job():
    """Daily 4 AM job: delete unaccepted invites older than 7 days."""
    with app.app_context():
        cutoff = datetime.utcnow() - timedelta(days=7)
        old = Invite.query.filter(
            Invite.accepted == False,  # noqa: E712
            Invite.created_at < cutoff,
        ).all()
        count = len(old)
        for invite in old:
            db.session.delete(invite)
        db.session.commit()
        log.info("Invite cleanup: deleted %d expired invite(s).", count)


scheduler = BackgroundScheduler(daemon=True)
scheduler.add_job(nightly_crawl_job, CronTrigger(hour=2, minute=0))
scheduler.add_job(nightly_cleanup_job, CronTrigger(hour=3, minute=0))
scheduler.add_job(nightly_widget_cleanup_job, CronTrigger(hour=3, minute=30))
scheduler.add_job(invite_cleanup_job, CronTrigger(hour=4, minute=0))
scheduler.add_job(trial_reminder_job, CronTrigger(hour=9, minute=0))
if not scheduler.running:
    scheduler.start()


# ── Billing helpers ───────────────────────────────────────────────────────────

# Email domains that are permanently exempt from billing checks.
# Add extra domains via the BILLING_EXEMPT_DOMAINS env var (comma-separated).
_extra_exempt  = {d.strip() for d in os.getenv("BILLING_EXEMPT_DOMAINS", "daltonfumc.com").split(",") if d.strip()}
EXEMPT_DOMAINS = {"wesleyai.co"} | _extra_exempt


def _is_billing_exempt(email: str) -> bool:
    domain = email.split("@")[-1].lower()
    return domain in EXEMPT_DOMAINS


def _require_active():
    """Return a redirect to /subscribe if the current church's billing has lapsed.
    Returns None if the user may continue.  Exempt domains and per-church
    billing_exempt flag always pass through.
    """
    if _is_billing_exempt(current_user.email):
        return None
    if current_user.church.billing_exempt:
        return None
    if not current_user.church.is_active:
        return redirect(url_for("subscribe_page"))
    return None


# ── Auth routes ───────────────────────────────────────────────────────────────


@app.route("/login")
def login_page():
    if current_user.is_authenticated:
        return redirect(url_for("chat_page"))
    return render_template("auth.html", mode="login")


@app.route("/signup")
def signup_page():
    if current_user.is_authenticated:
        return redirect(url_for("chat_page"))
    return render_template("auth.html", mode="signup")


@app.route("/api/auth/signup", methods=["POST"])
def api_signup():
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    password = (data.get("password") or "").strip()
    church_name = (data.get("church_name") or "").strip()

    if not email or not password or not church_name:
        return jsonify({"error": "Email, password, and church name are required."}), 400
    if len(password) < 8:
        return jsonify({"error": "Password must be at least 8 characters."}), 400
    if User.query.filter_by(email=email).first():
        return jsonify({"error": "An account with that email already exists."}), 400

    church = Church(
        name=church_name,
        trial_ends_at=datetime.utcnow() + timedelta(days=14),
    )
    db.session.add(church)
    db.session.flush()  # get church.id before commit

    user = User(
        email=email,
        password_hash=generate_password_hash(password, method="pbkdf2:sha256"),
        church_id=church.id,
    )
    db.session.add(user)
    db.session.commit()

    login_user(user)

    # Fire welcome email in background (non-blocking)
    _church_name     = church.name
    _trial_ends_at   = church.trial_ends_at
    threading.Thread(
        target=_send_welcome_email,
        args=(email, _church_name, _trial_ends_at),
        daemon=True,
    ).start()

    return jsonify({"ok": True}), 201


@app.route("/api/auth/login", methods=["POST"])
def api_login():
    data = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    password = (data.get("password") or "").strip()

    user = User.query.filter_by(email=email).first()
    if not user or not check_password_hash(user.password_hash, password):
        return jsonify({"error": "Invalid email or password."}), 401

    login_user(user)
    return jsonify({"ok": True})


@app.route("/logout")
def logout():
    logout_user()
    return redirect(url_for("login_page"))


# ── Password reset ────────────────────────────────────────────────────────────


def _send_reset_email(to_email: str, reset_url: str) -> None:
    """Send a branded HTML password-reset email via Resend."""
    html = render_template(
        "emails/reset_password.html",
        reset_url=reset_url,
        support_email=SUPPORT_EMAIL,
    )
    try:
        resend.Emails.send({
            "from": FROM_EMAIL,
            "to": [to_email],
            "subject": "Reset your Wesley AI password",
            "html": html,
        })
    except Exception as exc:
        log.error("Password reset email failed for %s: %s", to_email, exc)


def _send_welcome_email(to_email: str, church_name: str, trial_ends_at: datetime) -> None:
    """Send a branded welcome email to a new signup via Resend."""
    html = render_template(
        "emails/welcome.html",
        church_name=church_name,
        trial_date=trial_ends_at.strftime("%B %d, %Y"),
        app_url=APP_URL,
        support_email=SUPPORT_EMAIL,
    )
    try:
        resend.Emails.send({
            "from": FROM_EMAIL,
            "to": [to_email],
            "subject": f"Welcome to Wesley AI, {church_name}!",
            "html": html,
        })
    except Exception as exc:
        log.error("Welcome email failed for %s: %s", to_email, exc)


def _send_trial_expiring_email(to_email: str, church_name: str, trial_ends_at: datetime) -> None:
    """Send a trial-expiring warning email (4 days before trial ends) via Resend."""
    html = render_template(
        "emails/trial_expiring.html",
        church_name=church_name,
        trial_date=trial_ends_at.strftime("%B %d, %Y"),
        app_url=APP_URL,
        support_email=SUPPORT_EMAIL,
    )
    try:
        resend.Emails.send({
            "from": FROM_EMAIL,
            "to": [to_email],
            "subject": "Your Wesley AI trial ends in 4 days",
            "html": html,
        })
    except Exception as exc:
        log.error("Trial expiring email failed for %s: %s", to_email, exc)


def _send_payment_confirmation_email(to_email: str, church_name: str) -> None:
    """Send a payment confirmation email after a successful Stripe checkout via Resend."""
    html = render_template(
        "emails/payment_confirmation.html",
        church_name=church_name,
        app_url=APP_URL,
        support_email=SUPPORT_EMAIL,
    )
    try:
        resend.Emails.send({
            "from": FROM_EMAIL,
            "to": [to_email],
            "subject": "Your Wesley AI subscription is active",
            "html": html,
        })
    except Exception as exc:
        log.error("Payment confirmation email failed for %s: %s", to_email, exc)


def _send_invite_email(to_email: str, church_name: str, invite_url: str) -> None:
    """Send a branded staff invitation email via Resend."""
    html = render_template(
        "emails/invite.html",
        church_name=church_name,
        invite_url=invite_url,
        support_email=SUPPORT_EMAIL,
    )
    try:
        resend.Emails.send({
            "from": FROM_EMAIL,
            "to": [to_email],
            "subject": f"You've been invited to join {church_name} on Wesley AI",
            "html": html,
        })
    except Exception as exc:
        log.error("Invite email failed for %s: %s", to_email, exc)


@app.route("/forgot-password")
def forgot_password_page():
    if current_user.is_authenticated:
        return redirect(url_for("chat_page"))
    return render_template("forgot_password.html")


@app.route("/api/auth/forgot-password", methods=["POST"])
def api_forgot_password():
    data  = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()

    # Always return success to prevent email enumeration
    if not email:
        return jsonify({"ok": True})

    user = User.query.filter_by(email=email).first()
    if user:
        token = secrets.token_urlsafe(32)
        user.reset_token         = token
        user.reset_token_expires = datetime.utcnow() + timedelta(hours=1)
        db.session.commit()
        reset_url = url_for("reset_password_page", token=token, _external=True)
        _send_reset_email(user.email, reset_url)

    return jsonify({"ok": True})


@app.route("/reset-password/<token>")
def reset_password_page(token: str):
    if current_user.is_authenticated:
        return redirect(url_for("chat_page"))
    user = User.query.filter_by(reset_token=token).first()
    token_valid = (
        user is not None
        and user.reset_token_expires is not None
        and user.reset_token_expires > datetime.utcnow()
    )
    return render_template("reset_password.html", token=token, token_valid=token_valid)


@app.route("/api/auth/reset-password", methods=["POST"])
def api_reset_password():
    data     = request.get_json(silent=True) or {}
    token    = (data.get("token") or "").strip()
    password = (data.get("password") or "").strip()
    confirm  = (data.get("confirm") or "").strip()

    if not token or not password or not confirm:
        return jsonify({"error": "All fields are required."}), 400
    if password != confirm:
        return jsonify({"error": "Passwords do not match."}), 400
    if len(password) < 8:
        return jsonify({"error": "Password must be at least 8 characters."}), 400

    user = User.query.filter_by(reset_token=token).first()
    if not user or user.reset_token_expires is None or user.reset_token_expires <= datetime.utcnow():
        return jsonify({"error": "This reset link is invalid or has expired."}), 400

    user.password_hash       = generate_password_hash(password, method="pbkdf2:sha256")
    user.reset_token         = None
    user.reset_token_expires = None
    db.session.commit()

    return jsonify({"ok": True})


# ── Chat page (main interface) ────────────────────────────────────────────────


@app.route("/")
@login_required
def chat_page():
    church = current_user.church
    if not church.onboarding_complete:
        return redirect(url_for("onboarding_page"))
    check = _require_active()
    if check:
        return check
    branding = _build_branding_dict(church)
    return render_template(
        "dashboard.html",
        church_name=church.name,
        user_email=current_user.email,
        bot_name=branding["bot_name"],
        welcome_message=branding["welcome_message"],
        primary_color=branding["primary_color"],
        starter_questions=json.dumps(branding["starter_questions"]),
    )


# ── Onboarding wizard ─────────────────────────────────────────────────────────


@app.route("/onboarding")
@login_required
def onboarding_page():
    if current_user.church.onboarding_complete:
        return redirect(url_for("chat_page"))
    church = current_user.church
    return render_template(
        "onboarding.html",
        church_name=church.name,
        church_id=church.id,
    )


@app.route("/api/onboarding/step1", methods=["POST"])
@login_required
def onboarding_step1():
    """Save church name + city and mark onboarding complete."""
    data = request.get_json(silent=True) or {}
    church_name = (data.get("church_name") or "").strip()
    church_city = (data.get("church_city") or "").strip()

    if not church_name:
        return jsonify({"error": "Church name cannot be empty."}), 400

    church = current_user.church
    church.name = church_name[:200]
    church.church_city = church_city[:200] if church_city else None
    church.onboarding_complete = True
    db.session.commit()
    return jsonify({"ok": True})


# ── Dashboard (management) ────────────────────────────────────────────────────


@app.route("/dashboard")
@login_required
def management_dashboard():
    # Staff members can only access the chat page
    if current_user.role == "staff":
        return redirect(url_for("chat_page"))
    check = _require_active()
    if check:
        return check
    church = current_user.church
    branding = _build_branding_dict(church)
    return render_template(
        "settings.html",
        church_name=church.name,
        church_id=current_user.church_id,
        user_email=current_user.email,
        user_role=current_user.role,
        bot_name=branding["bot_name"],
        welcome_message=branding["welcome_message"],
        primary_color=branding["primary_color"],
        church_city=branding["church_city"],
        has_stripe_sub=bool(church.stripe_subscription_id),
    )


# ── Staff management API ──────────────────────────────────────────────────────


@app.route("/api/staff")
@login_required
def list_staff():
    """Return all users belonging to the current church (admin only)."""
    if current_user.role != "admin":
        return jsonify({"error": "Forbidden."}), 403
    users = (
        User.query
        .filter_by(church_id=current_user.church_id)
        .order_by(User.created_at)
        .all()
    )
    # Include pending (unaccepted) invites so the UI can show them too
    pending = (
        Invite.query
        .filter_by(church_id=current_user.church_id, accepted=False)
        .order_by(Invite.created_at)
        .all()
    )
    return jsonify({
        "staff": [
            {"id": u.id, "email": u.email, "role": u.role, "created_at": u.created_at.isoformat()}
            for u in users
        ],
        "pending_invites": [
            {"id": inv.id, "email": inv.email, "created_at": inv.created_at.isoformat()}
            for inv in pending
        ],
    })


@app.route("/api/staff/invite", methods=["POST"])
@login_required
def invite_staff():
    """Send a staff invitation email (admin only)."""
    if current_user.role != "admin":
        return jsonify({"error": "Forbidden."}), 403

    data  = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()

    if not email:
        return jsonify({"error": "Email is required."}), 400

    # Don't invite someone who already has an account at this church
    existing = User.query.filter_by(email=email, church_id=current_user.church_id).first()
    if existing:
        return jsonify({"error": "A user with that email already exists on your team."}), 400

    # Don't create a duplicate pending invite
    dup = Invite.query.filter_by(
        email=email, church_id=current_user.church_id, accepted=False
    ).first()
    if dup:
        return jsonify({"error": "An invitation has already been sent to that email."}), 400

    token  = secrets.token_urlsafe(32)
    invite = Invite(
        church_id=current_user.church_id,
        email=email,
        token=token,
    )
    db.session.add(invite)
    db.session.commit()

    invite_url = url_for("accept_invite_page", token=token, _external=True)
    church_name = current_user.church.name
    threading.Thread(
        target=_send_invite_email,
        args=(email, church_name, invite_url),
        daemon=True,
    ).start()

    return jsonify({"ok": True}), 201


@app.route("/api/staff/<int:user_id>", methods=["DELETE"])
@login_required
def remove_staff(user_id):
    """Remove a staff user from the church (admin only; cannot remove admins or self)."""
    if current_user.role != "admin":
        return jsonify({"error": "Forbidden."}), 403

    if user_id == current_user.id:
        return jsonify({"error": "You cannot remove yourself."}), 400

    user = User.query.filter_by(id=user_id, church_id=current_user.church_id).first()
    if not user:
        return jsonify({"error": "User not found."}), 404

    if user.role == "admin":
        return jsonify({"error": "Admin accounts cannot be removed via this endpoint."}), 400

    db.session.delete(user)
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/invite/<token>")
def accept_invite_page(token: str):
    """Public invite acceptance page — validates token and renders invite.html."""
    invite = Invite.query.filter_by(token=token, accepted=False).first()
    cutoff = datetime.utcnow() - timedelta(days=7)
    token_valid = (
        invite is not None
        and invite.created_at >= cutoff
    )
    church_name = ""
    if token_valid and invite:
        church = Church.query.get(invite.church_id)
        church_name = church.name if church else ""
    return render_template(
        "invite.html",
        token=token,
        token_valid=token_valid,
        church_name=church_name,
    )


@app.route("/api/invite/accept", methods=["POST"])
def api_accept_invite():
    """Create a staff account from a valid invite token."""
    data     = request.get_json(silent=True) or {}
    token    = (data.get("token") or "").strip()
    password = (data.get("password") or "").strip()
    confirm  = (data.get("confirm") or "").strip()

    if not token or not password or not confirm:
        return jsonify({"error": "All fields are required."}), 400
    if password != confirm:
        return jsonify({"error": "Passwords do not match."}), 400
    if len(password) < 8:
        return jsonify({"error": "Password must be at least 8 characters."}), 400

    invite = Invite.query.filter_by(token=token, accepted=False).first()
    cutoff = datetime.utcnow() - timedelta(days=7)
    if not invite or invite.created_at < cutoff:
        return jsonify({"error": "This invitation link is invalid or has expired."}), 400

    # Guard: email may already have an account (e.g. double-click)
    if User.query.filter_by(email=invite.email).first():
        return jsonify({"error": "An account with this email already exists. Please log in."}), 400

    user = User(
        email=invite.email,
        password_hash=generate_password_hash(password, method="pbkdf2:sha256"),
        church_id=invite.church_id,
        role="staff",
    )
    db.session.add(user)
    invite.accepted = True
    db.session.commit()

    login_user(user)
    return jsonify({"ok": True})


# ── Documents API ─────────────────────────────────────────────────────────────


@app.route("/api/documents")
@login_required
def list_documents():
    docs = (
        Document.query
        .filter_by(church_id=current_user.church_id)
        .order_by(Document.uploaded_at.desc())
        .all()
    )
    return jsonify({
        "documents": [
            {
                "id": d.id,
                "name": d.original_name,
                "size_kb": round(d.size_bytes / 1024, 1),
                "type": Path(d.original_name).suffix.lower().lstrip("."),
                "visibility": d.visibility,
            }
            for d in docs
        ]
    })


@app.route("/api/documents/upload", methods=["POST"])
@login_required
def upload_document():
    if "file" not in request.files:
        return jsonify({"error": "No file provided."}), 400
    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "No file selected."}), 400

    original_name = file.filename
    suffix = Path(original_name).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        return jsonify({"error": "Only PDF and DOCX files are supported."}), 400

    content = file.read()
    size_bytes = len(content)

    # UUID filename prevents path traversal and collisions
    stored_name = f"{uuid.uuid4().hex}{suffix}"
    church_dir = get_church_dir(current_user.church_id)
    (church_dir / stored_name).write_bytes(content)

    # sanitize display name; fall back to stored name if result is empty
    display_name = secure_filename(original_name) or stored_name

    visibility = request.form.get("visibility", "staff_only")
    if visibility not in ("staff_only", "staff_and_chatbot"):
        visibility = "staff_only"

    doc = Document(
        church_id=current_user.church_id,
        filename=stored_name,
        original_name=display_name,
        size_bytes=size_bytes,
        visibility=visibility,
    )
    db.session.add(doc)
    db.session.commit()

    return jsonify({
        "ok": True,
        "id": doc.id,
        "name": doc.original_name,
        "size_kb": round(size_bytes / 1024, 1),
        "type": suffix.lstrip("."),
        "visibility": doc.visibility,
    }), 201


@app.route("/api/documents/<int:doc_id>", methods=["DELETE"])
@login_required
def delete_document(doc_id):
    doc = Document.query.filter_by(id=doc_id, church_id=current_user.church_id).first()
    if not doc:
        return jsonify({"error": "Document not found."}), 404

    filepath = get_church_dir(current_user.church_id) / doc.filename
    if filepath.exists():
        filepath.unlink()

    _evict_doc_cache(doc.id, doc.uploaded_at)
    db.session.delete(doc)
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/documents/<int:doc_id>", methods=["PATCH"])
@login_required
def update_document_visibility(doc_id):
    doc = Document.query.filter_by(id=doc_id, church_id=current_user.church_id).first()
    if not doc:
        return jsonify({"error": "Document not found."}), 404

    data = request.get_json(silent=True) or {}
    visibility = data.get("visibility")
    if visibility not in ("staff_only", "staff_and_chatbot"):
        return jsonify({"error": "Invalid visibility value."}), 400

    doc.visibility = visibility
    db.session.commit()
    return jsonify({"ok": True, "visibility": doc.visibility})


# ── Chat API (staff dashboard) ────────────────────────────────────────────────


@app.route("/api/chat", methods=["POST"])
@login_required
def chat():
    data = request.get_json(silent=True)
    if not data or not data.get("question", "").strip():
        return jsonify({"error": "No question provided"}), 400

    question = data["question"].strip()
    conversation_id = data.get("conversation_id")

    # Resolve or create the conversation
    if conversation_id:
        conv = Conversation.query.filter_by(
            id=conversation_id, church_id=current_user.church_id
        ).first()
        if not conv:
            return jsonify({"error": "Conversation not found."}), 404
    else:
        title = question[:40]
        conv = Conversation(church_id=current_user.church_id, title=title)
        db.session.add(conv)
        db.session.flush()  # get conv.id

    # Build history BEFORE adding the new user message to avoid the autoflush
    # pitfall: accessing conv.messages triggers a SELECT, which autoflush would
    # include the just-added user message in — duplicating the question as two
    # consecutive user turns and causing Gemini to reject the request.
    history = [
        {"role": m.role, "content": m.content}
        for m in conv.messages
    ]

    # Save the user message (after history snapshot so it's not duplicated)
    db.session.add(Message(conversation_id=conv.id, role="user", content=question))

    # RAG context
    chunks = load_church_documents(current_user.church_id)
    context = ""
    candidate_sources = []

    if chunks:
        scored = find_relevant_chunks(question, chunks)
        if scored:
            context = build_context_block(scored)
            seen: set[tuple] = set()
            for _, chunk in scored:
                key = (chunk["source"], chunk["location"])
                if key not in seen:
                    seen.add(key)
                    candidate_sources.append({"file": chunk["source"], "location": chunk["location"]})

    system_instruction = _build_system_prompt(current_user.church, widget=False)

    try:
        answer = call_gemini(question, context, history, system_instruction)
    except ValueError as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500
    except Exception as e:
        db.session.rollback()
        user_msg, status = _friendly_gemini_error(e)
        return jsonify({"error": user_msg}), status

    # Save the assistant message and touch updated_at
    db.session.add(Message(conversation_id=conv.id, role="assistant", content=answer))
    conv.updated_at = datetime.utcnow()
    db.session.commit()

    answer_lower = answer.lower()
    sources = [s for s in candidate_sources if s["file"].lower() in answer_lower]

    return jsonify({"answer": answer, "sources": sources, "conversation_id": conv.id})


# ── Conversations API ──────────────────────────────────────────────────────────


@app.route("/api/conversations")
@login_required
def list_conversations():
    convs = (
        Conversation.query
        .filter_by(church_id=current_user.church_id)
        .order_by(Conversation.updated_at.desc())
        .all()
    )
    return jsonify({
        "conversations": [
            {"id": c.id, "title": c.title, "updated_at": c.updated_at.isoformat()}
            for c in convs
        ]
    })


@app.route("/api/conversations/<int:conv_id>/messages")
@login_required
def get_conversation_messages(conv_id):
    conv = Conversation.query.filter_by(
        id=conv_id, church_id=current_user.church_id
    ).first()
    if not conv:
        return jsonify({"error": "Conversation not found."}), 404
    return jsonify({
        "conversation_id": conv.id,
        "title": conv.title,
        "messages": [
            {"role": m.role, "content": m.content, "created_at": m.created_at.isoformat()}
            for m in conv.messages
        ],
    })


# ── Widget Conversations API (staff dashboard) ────────────────────────────────


@app.route("/api/widget/conversations")
@login_required
def list_widget_conversations():
    wconvs = (
        WidgetConversation.query
        .options(joinedload(WidgetConversation.messages))
        .filter_by(church_id=current_user.church_id)
        .order_by(WidgetConversation.updated_at.desc())
        .all()
    )
    result = []
    for wc in wconvs:
        # First user message as a preview
        first_msg = next((m for m in wc.messages if m.role == "user"), None)
        preview = (first_msg.content[:80] + "…") if first_msg and len(first_msg.content) > 80 else (first_msg.content if first_msg else "")
        result.append({
            "id": wc.id,
            "session_id": wc.session_id,
            "created_at": wc.created_at.isoformat(),
            "updated_at": wc.updated_at.isoformat(),
            "preview": preview,
            "message_count": len(wc.messages),
        })
    return jsonify({"conversations": result})


@app.route("/api/widget/conversations/<int:wconv_id>/messages")
@login_required
def get_widget_conversation_messages(wconv_id):
    wconv = WidgetConversation.query.filter_by(
        id=wconv_id, church_id=current_user.church_id
    ).first()
    if not wconv:
        return jsonify({"error": "Widget conversation not found."}), 404
    return jsonify({
        "id": wconv.id,
        "session_id": wconv.session_id,
        "messages": [
            {"role": m.role, "content": m.content, "created_at": m.created_at.isoformat()}
            for m in wconv.messages
        ],
    })


# ── Church branding API ───────────────────────────────────────────────────────

_HEX_COLOR_RE = re.compile(r"^#[0-9a-fA-F]{6}$")


@app.route("/api/church/branding", methods=["GET"])
@login_required
def get_church_branding():
    return jsonify(_build_branding_dict(current_user.church))


@app.route("/api/church/branding", methods=["POST"])
@login_required
def save_church_branding():
    data = request.get_json(silent=True) or {}
    church = current_user.church

    bot_name = (data.get("bot_name") or "").strip()
    bot_subtitle = (data.get("bot_subtitle") or "").strip()
    welcome_message = (data.get("welcome_message") or "").strip()
    primary_color = (data.get("primary_color") or "").strip()
    church_city = (data.get("church_city") or "").strip()
    raw_sugs = data.get("starter_questions") or []

    if not bot_name:
        return jsonify({"error": "Bot name cannot be empty."}), 400
    if not welcome_message:
        return jsonify({"error": "Welcome message cannot be empty."}), 400
    if primary_color and not _HEX_COLOR_RE.match(primary_color):
        return jsonify({"error": "Primary color must be a valid hex color (e.g. #1a2b3c)."}), 400

    # Sanitise starter questions: keep only non-empty strings, max 4, max 200 chars each
    clean_sugs = [str(s).strip()[:200] for s in raw_sugs if str(s).strip()][:4]

    church.bot_name = bot_name[:100]
    church.bot_subtitle = bot_subtitle[:200] if bot_subtitle else None
    church.welcome_message = welcome_message[:500]
    church.primary_color = primary_color if primary_color else DEFAULT_COLOR
    church.church_city = church_city[:200] if church_city else None
    church.starter_questions = json.dumps(clean_sugs) if clean_sugs else None
    db.session.commit()
    return jsonify({"ok": True})


# ── Admin panel ───────────────────────────────────────────────────────────────


@app.route("/admin")
@login_required
def admin_panel():
    if not is_super_admin():
        return render_template("admin.html", forbidden=True), 403
    prompt_row = SystemPrompt.query.get(1)
    current_prompt = prompt_row.content if prompt_row else DEFAULT_SYSTEM_PROMPT
    return render_template(
        "admin.html",
        forbidden=False,
        current_prompt=current_prompt,
        default_prompt=DEFAULT_SYSTEM_PROMPT,
    )


@app.route("/api/admin/system-prompt", methods=["POST"])
@login_required
def update_system_prompt():
    if not is_super_admin():
        return jsonify({"error": "Forbidden."}), 403
    data = request.get_json(silent=True) or {}
    content = (data.get("content") or "").strip()
    if not content:
        return jsonify({"error": "Prompt content cannot be empty."}), 400
    prompt_row = SystemPrompt.query.get(1)
    if prompt_row:
        prompt_row.content = content
    else:
        db.session.add(SystemPrompt(id=1, content=content))
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/admin/churches", methods=["GET"])
@login_required
def admin_list_churches():
    if not is_super_admin():
        return jsonify({"error": "Forbidden."}), 403

    churches = Church.query.order_by(Church.created_at.desc()).all()

    total_messages = db.session.execute(
        text("SELECT COUNT(*) FROM messages")
    ).scalar() or 0
    total_widget_messages = db.session.execute(
        text("SELECT COUNT(*) FROM widget_messages")
    ).scalar() or 0
    total_all_messages = total_messages + total_widget_messages

    now = datetime.utcnow()
    active_subs = 0
    trialing = 0
    for c in churches:
        if c.stripe_subscription_id:
            active_subs += 1
        elif c.trial_ends_at and c.trial_ends_at > now:
            trialing += 1

    stats = {
        "total_churches":       len(churches),
        "total_messages":       total_all_messages,
        "active_subscriptions": active_subs,
        "trialing":             trialing,
    }

    church_list = []
    for c in churches:
        # admin email = first user
        first_user = User.query.filter_by(church_id=c.id).order_by(User.created_at).first()
        admin_email = first_user.email if first_user else ""

        msg_count = db.session.execute(
            text("SELECT COUNT(*) FROM messages m "
                 "JOIN conversations cv ON cv.id = m.conversation_id "
                 "WHERE cv.church_id = :cid"),
            {"cid": c.id}
        ).scalar() or 0

        widget_msg_count = db.session.execute(
            text("SELECT COUNT(*) FROM widget_messages wm "
                 "JOIN widget_conversations wc ON wc.id = wm.widget_conversation_id "
                 "WHERE wc.church_id = :cid"),
            {"cid": c.id}
        ).scalar() or 0

        doc_count = db.session.execute(
            text("SELECT COUNT(*) FROM documents WHERE church_id = :cid"),
            {"cid": c.id}
        ).scalar() or 0

        if c.billing_exempt:
            status = "exempt"
        elif c.stripe_subscription_id:
            status = "active"
        elif c.trial_ends_at and c.trial_ends_at > now:
            status = "trialing"
        else:
            status = "expired"

        church_list.append({
            "id":                    c.id,
            "name":                  c.name,
            "admin_email":           admin_email,
            "church_city":           c.church_city or "",
            "created_at":            c.created_at.isoformat() if c.created_at else "",
            "trial_ends_at":         c.trial_ends_at.isoformat() if c.trial_ends_at else "",
            "plan":                  c.plan or "founders",
            "billing_exempt":        c.billing_exempt,
            "stripe_subscription_id": c.stripe_subscription_id or "",
            "status":                status,
            "message_count":         msg_count,
            "widget_message_count":  widget_msg_count,
            "doc_count":             doc_count,
        })

    return jsonify({"stats": stats, "churches": church_list})


@app.route("/api/admin/churches/<int:church_id>", methods=["PATCH"])
@login_required
def admin_update_church(church_id):
    if not is_super_admin():
        return jsonify({"error": "Forbidden."}), 403

    church = Church.query.get_or_404(church_id)
    data = request.get_json(silent=True) or {}

    if "name" in data:
        name = data["name"].strip()
        if name:
            church.name = name

    if "church_city" in data:
        church.church_city = data["church_city"].strip() or None

    if "plan" in data:
        plan = data["plan"].strip()
        if plan in ("founders", "small", "medium", "large"):
            church.plan = plan

    if "trial_ends_at" in data:
        val = data["trial_ends_at"]
        if not val:
            church.trial_ends_at = None
        else:
            try:
                # Accept ISO date (YYYY-MM-DD) or full ISO datetime
                if "T" in str(val):
                    church.trial_ends_at = datetime.fromisoformat(str(val)[:19])
                else:
                    church.trial_ends_at = datetime.strptime(str(val)[:10], "%Y-%m-%d")
            except ValueError:
                return jsonify({"error": "Invalid trial_ends_at format. Use YYYY-MM-DD."}), 400

    if "billing_exempt" in data:
        church.billing_exempt = bool(data["billing_exempt"])

    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/admin/churches/<int:church_id>/reset-password", methods=["POST"])
@login_required
def admin_reset_password(church_id):
    if not is_super_admin():
        return jsonify({"error": "Forbidden."}), 403

    church = Church.query.get_or_404(church_id)
    user = User.query.filter_by(church_id=church.id).order_by(User.created_at).first()
    if not user:
        return jsonify({"error": "No user found for this church."}), 404

    alphabet = string.ascii_letters + string.digits
    temp_password = "".join(secrets.choice(alphabet) for _ in range(12))
    user.password_hash = generate_password_hash(temp_password)
    db.session.commit()

    return jsonify({"email": user.email, "temp_password": temp_password})


# ── Widget JS (public, CORS) ──────────────────────────────────────────────────


def _widget_js_response():
    """Serve widget-core.js with appropriate headers for public embedding."""
    resp = make_response(send_from_directory("static", "widget-core.js"))
    resp.headers["Content-Type"] = "application/javascript; charset=utf-8"
    resp.headers["Access-Control-Allow-Origin"] = "*"
    # no-cache so church sites always validate against the server and get
    # the latest widget code immediately without needing a cache-busting URL.
    resp.headers["Cache-Control"] = "no-cache, must-revalidate"
    return resp


@app.route("/widget.js")
def serve_widget():
    return _widget_js_response()


@app.route("/widget-core.js")
def serve_widget_core():
    return _widget_js_response()


# ── Widget Branding API (public, CORS) ────────────────────────────────────────


@app.route("/api/widget/branding", methods=["GET", "OPTIONS"])
def widget_branding():
    """Public endpoint returning church branding config for embedded widgets."""
    if request.method == "OPTIONS":
        resp = make_response("", 204)
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
        return resp

    if _widget_branding_limiter.is_limited(request.remote_addr or "unknown"):
        resp = jsonify({"error": "Rate limit exceeded. Please try again later."})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, 429

    church_id = request.args.get("church_id", "").strip()
    if not church_id:
        return jsonify({"error": "church_id is required"}), 400

    try:
        church_id_int = int(church_id)
    except ValueError:
        return jsonify({"error": "Invalid church_id"}), 400

    church = Church.query.get(church_id_int)
    if not church:
        return jsonify({"error": "Church not found"}), 404

    resp = jsonify(_build_branding_dict(church))
    resp.headers["Access-Control-Allow-Origin"] = "*"
    resp.headers["Cache-Control"] = "public, max-age=60"
    return resp


# ── Widget Chat API (public, CORS, crawled content only) ─────────────────────


@app.route("/api/widget/chat", methods=["POST", "OPTIONS"])
def widget_chat():
    # Handle CORS preflight
    if request.method == "OPTIONS":
        resp = make_response("", 204)
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
        return resp

    if _widget_chat_limiter.is_limited(request.remote_addr or "unknown"):
        resp = jsonify({"error": "Rate limit exceeded. Please try again later."})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, 429

    data = request.get_json(silent=True) or {}
    church_id_raw = data.get("church_id")
    question = (data.get("question") or "").strip()
    session_id = (data.get("session_id") or "").strip() or None

    def cors_err(msg, status=400):
        resp = jsonify({"error": msg})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, status

    if not church_id_raw or not question:
        return cors_err("church_id and question are required.")

    try:
        church_id = int(church_id_raw)
    except (ValueError, TypeError):
        return cors_err("Invalid church_id.")

    church = Church.query.get(church_id)
    if not church:
        return cors_err("Church not found.", 404)

    # Resolve or create the widget conversation for this session
    wconv = None
    if session_id:
        wconv = WidgetConversation.query.filter_by(
            church_id=church_id, session_id=session_id
        ).first()
    if not wconv:
        session_id = uuid.uuid4().hex
        wconv = WidgetConversation(church_id=church_id, session_id=session_id)
        db.session.add(wconv)
        db.session.flush()  # get wconv.id

    # Build history BEFORE adding the new user message.
    # SQLAlchemy autoflush fires before any SELECT (including lazy relationship
    # loads), so if we added the user message first, it would be included in
    # wconv.messages and appear twice in the Gemini content list — causing a
    # "consecutive user turns" error that rolls back the whole transaction.
    history = [
        {"role": m.role, "content": m.content}
        for m in wconv.messages
    ]

    # Save the user message (after history snapshot so it's not duplicated)
    db.session.add(WidgetMessage(
        widget_conversation_id=wconv.id, role="user", content=question
    ))

    # ── Build RAG context with document priority ───────────────────────────
    # Documents marked staff_and_chatbot are ALWAYS included when they exist —
    # they bypass the score > 0 threshold because the church explicitly enabled
    # them for the widget. They are sorted by relevance but never filtered out.
    # Web chunks fill any remaining slots using the normal scoring filter.
    MAX_DOC_CHUNKS = 5
    MAX_WEB_CHUNKS = 5

    web_chunks = load_church_web_content(church_id)
    doc_chunks = load_chatbot_documents(church_id)

    scored_docs: list[tuple[int, dict]] = []
    if doc_chunks:
        keywords = extract_keywords(question)
        if keywords:
            scored_docs = sorted(
                [(score_chunk(c, keywords), c) for c in doc_chunks],
                key=lambda x: x[0], reverse=True,
            )[:MAX_DOC_CHUNKS]
        else:
            # No keywords extracted — include the first N doc chunks as-is
            scored_docs = [(0, c) for c in doc_chunks[:MAX_DOC_CHUNKS]]

    scored_web = find_relevant_chunks(question, web_chunks, top_n=MAX_WEB_CHUNKS) if web_chunks else []

    # Documents go first so Gemini sees church-specific content before web copy
    context_parts = []
    if scored_docs:
        context_parts.append(build_context_block(scored_docs))
    if scored_web:
        context_parts.append(build_context_block(scored_web))
    context = "\n".join(context_parts)

    system_instruction = _build_system_prompt(church, widget=True)

    try:
        answer = call_gemini(question, context, history, system_instruction)
    except ValueError as e:
        db.session.rollback()
        return cors_err(str(e), 500)
    except Exception as e:
        db.session.rollback()
        user_msg, status = _friendly_gemini_error(e)
        return cors_err(user_msg, status)

    # Save the assistant message and touch updated_at
    db.session.add(WidgetMessage(
        widget_conversation_id=wconv.id, role="assistant", content=answer
    ))
    wconv.updated_at = datetime.utcnow()
    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        log.error("[WIDGET] DB commit failed: %s", e)
        return cors_err("Failed to save conversation. Please try again.", 500)

    resp = jsonify({"answer": answer, "session_id": session_id})
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp


# ── Church settings API (website URL + crawl stats) ───────────────────────────


@app.route("/api/church/settings", methods=["GET"])
@login_required
def get_church_settings():
    church = current_user.church
    page_count = CrawledPage.query.filter_by(church_id=church.id).count()
    return jsonify({
        "website_url": church.website_url or "",
        "last_crawled_at": church.last_crawled_at.isoformat() if church.last_crawled_at else None,
        "page_count": page_count,
        "church_id": church.id,
    })


@app.route("/api/church/settings", methods=["POST"])
@login_required
def save_church_settings():
    data = request.get_json(silent=True) or {}
    url = (data.get("website_url") or "").strip().rstrip("/")
    if url and not url.startswith(("http://", "https://")):
        return jsonify({"error": "URL must start with http:// or https://"}), 400
    current_user.church.website_url = url or None
    db.session.commit()
    return jsonify({"ok": True})


# ── Manual re-crawl (fires background thread) ─────────────────────────────────


@app.route("/api/church/crawl", methods=["POST"])
@login_required
def trigger_crawl():
    church = current_user.church
    if not church.website_url:
        return jsonify({"error": "No website URL configured. Save a URL first."}), 400

    crawl_url  = church.website_url
    church_id  = church.id

    def run_crawl():
        with app.app_context():
            from crawler import crawl_church_website
            result = crawl_church_website(church_id, crawl_url)
            log.info("Manual crawl church_id=%d: %s", church_id, result)

    t = threading.Thread(target=run_crawl, daemon=True)
    t.start()

    return jsonify({"ok": True, "message": "Crawl started in the background."})


# ── Stripe billing ────────────────────────────────────────────────────────────


@app.route("/subscribe")
@login_required
def subscribe_page():
    church = current_user.church
    days_left = None
    if church.trial_ends_at and church.trial_ends_at > datetime.utcnow():
        days_left = (church.trial_ends_at - datetime.utcnow()).days
    return render_template("subscribe.html",
                           user_email=current_user.email,
                           days_left=days_left)


@app.route("/stripe/checkout", methods=["POST"])
@login_required
def stripe_checkout():
    _validate_csrf()
    if not stripe.api_key:
        return "STRIPE_SECRET_KEY is not configured.", 500

    # Choose price based on billing cycle (monthly is default)
    billing_cycle = request.form.get("billing_cycle", "monthly")
    if billing_cycle == "annual":
        price_id = os.getenv("STRIPE_ANNUAL_PRICE_ID")
        if not price_id:
            return "STRIPE_ANNUAL_PRICE_ID is not configured.", 500
    else:
        price_id = os.getenv("STRIPE_MONTHLY_PRICE_ID")
        if not price_id:
            return "STRIPE_MONTHLY_PRICE_ID is not configured.", 500

    church = current_user.church

    try:
        # Retrieve or create a Stripe Customer so the portal works later
        customer_id = church.stripe_customer_id
        if not customer_id:
            existing = stripe.Customer.list(email=current_user.email, limit=1)
            if existing.data:
                customer_id = existing.data[0].id
            else:
                customer = stripe.Customer.create(
                    email=current_user.email,
                    metadata={"church_id": str(current_user.church_id)},
                )
                customer_id = customer.id
            church.stripe_customer_id = customer_id
            db.session.commit()

        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            mode="subscription",
            customer=customer_id,
            client_reference_id=str(current_user.church_id),
            line_items=[{"price": price_id, "quantity": 1}],
            success_url=url_for("stripe_success", _external=True),
            cancel_url=url_for("subscribe_page", _external=True),
        )
        return redirect(session.url, code=303)
    except stripe.StripeError as e:
        return render_template("subscribe.html",
                               user_email=current_user.email,
                               days_left=None,
                               stripe_error=getattr(e, "user_message", str(e)))


@app.route("/stripe/success")
@login_required
def stripe_success():
    return render_template("stripe_success.html")


@app.route("/billing/portal")
@login_required
def billing_portal():
    """Redirect the logged-in church admin to the Stripe Customer Portal."""
    if not stripe.api_key:
        return "STRIPE_SECRET_KEY is not configured.", 500

    church = current_user.church
    customer_id = church.stripe_customer_id

    # Fallback: look up the Stripe customer by email if we don't have the ID stored
    if not customer_id:
        try:
            existing = stripe.Customer.list(email=current_user.email, limit=1)
            if existing.data:
                customer_id = existing.data[0].id
                church.stripe_customer_id = customer_id
                db.session.commit()
        except stripe.StripeError:
            pass

    if not customer_id:
        return render_template(
            "subscribe.html",
            user_email=current_user.email,
            days_left=None,
            stripe_error="No billing account found. Please subscribe first.",
        )

    try:
        portal_session = stripe.billing_portal.Session.create(
            customer=customer_id,
            return_url=url_for("management_dashboard", _external=True),
        )
        return redirect(portal_session.url, code=303)
    except stripe.StripeError as e:
        return render_template(
            "subscribe.html",
            user_email=current_user.email,
            days_left=None,
            stripe_error=getattr(e, "user_message", str(e)),
        )


@app.route("/stripe/webhook", methods=["POST"])
def stripe_webhook():
    payload       = request.get_data()
    sig_header    = request.headers.get("Stripe-Signature", "")
    webhook_secret = os.getenv("STRIPE_WEBHOOK_SECRET", "")

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
    except ValueError:
        return jsonify({"error": "Invalid payload"}), 400
    except stripe.SignatureVerificationError:
        return jsonify({"error": "Invalid signature"}), 400

    etype = event["type"]

    if etype == "checkout.session.completed":
        sess        = event["data"]["object"]
        church_ref  = sess.get("client_reference_id")
        sub_id      = sess.get("subscription")
        customer_id = sess.get("customer")
        if church_ref and sub_id:
            try:
                church_id_int = int(church_ref)
            except (ValueError, TypeError):
                log.error("Stripe webhook: invalid client_reference_id=%r", church_ref)
                return jsonify({"ok": True})
            church = Church.query.get(church_id_int)
            if church:
                church.stripe_subscription_id = sub_id
                if customer_id and not church.stripe_customer_id:
                    church.stripe_customer_id = customer_id
                db.session.commit()
                log.info("Stripe webhook: church_id=%s subscribed (%s)", church_ref, sub_id)

                # Send payment confirmation email (get billing contact from first user)
                first_user = User.query.filter_by(church_id=church.id).order_by(User.id).first()
                if first_user:
                    _cname = church.name
                    _email = first_user.email
                    threading.Thread(
                        target=_send_payment_confirmation_email,
                        args=(_email, _cname),
                        daemon=True,
                    ).start()

    elif etype == "customer.subscription.deleted":
        sub    = event["data"]["object"]
        sub_id = sub["id"]
        church = Church.query.filter_by(stripe_subscription_id=sub_id).first()
        if church:
            church.stripe_subscription_id = None
            db.session.commit()
            log.info("Stripe webhook: subscription %s cancelled", sub_id)

    return jsonify({"ok": True})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5001, debug=os.getenv("FLASK_DEBUG", "false").lower() in ("1", "true"))
